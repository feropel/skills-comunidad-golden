"""
generate_report.py — Genera el documento Word del análisis de Meta Ads.

Reemplaza la versión .js anterior. Usa python-docx (preinstalado en el sandbox)
para eliminar la dependencia de Node y la fragilidad de require('docx').

Uso programático:
    from generate_report import generate_word_report
    path = generate_word_report(config_dict, output_path)

Uso CLI:
    python generate_report.py <config.json> <output.docx>

Estructura del config_dict (todas las claves son opcionales — el reporte
omite las secciones para las que falte data):

{
    "producto": "Sérum 7 Days",
    "periodo": "25 abril - 10 mayo 2026",
    "moneda_original": "COP",         # o "USD"
    "trm_aplicada": None,             # solo si hubo conversión USD→COP
    "trm_fuente": None,               # "web_search" | "fallback_4000" | "usuario"
    "ue_inputs": {...},               # ticket, costo, flete, tasas...
    "unit_economics": {...},          # output de calculate_unit_economics
    "pyl_historico": {...},           # output de calcular_pyl_historico
    "totales": {"gasto_cop": 37452127, "compras": 1109, "cpa_promedio": 33771},
    "campañas_ranking": [...],
    "creativos_ranking": [...],
    "demografia_tabla": [...],
    "ubicaciones_tabla": [...],
    "landings_tabla": [...],
    "embudo": {...},
    "plan_accion": {
        "pausar": [...],
        "escalar": [...],
        "campañas_nuevas": [...],
        "replicar_mismo_bm": [...],
        "replicar_bm_nuevo": [...]
    },
    "sintesis_ejecutiva": [...],      # 7-10 puntos
    "produccion_creativa": {...},
}
"""

import json
import sys
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# COLORES CORPORATIVOS GOLDEN GROUP
# ============================================================
COLORS = {
    "primary":  RGBColor(0x1F, 0x4E, 0x79),  # Azul corporativo
    "accent":   RGBColor(0xC0, 0x00, 0x00),  # Rojo (alertas)
    "green":    RGBColor(0x54, 0x82, 0x35),  # Verde (ganadores)
    "amber":    RGBColor(0xBF, 0x8F, 0x00),  # Amarillo (revisar)
    "gray":     RGBColor(0x59, 0x59, 0x59),
    "white":    RGBColor(0xFF, 0xFF, 0xFF),
    "black":    RGBColor(0x00, 0x00, 0x00),
}

SHADE = {
    "primary":  "1F4E79",
    "light":    "E7EEF5",
    "lightG":   "E2EFDA",
    "lightR":   "FCE4D6",
    "lightY":   "FFF2CC",
    "border":   "BFBFBF",
}


# ============================================================
# HELPERS
# ============================================================

def _cop(value):
    """Formatea valor numérico como pesos colombianos."""
    if value is None:
        return "—"
    try:
        return f"${float(value):,.0f}"
    except (TypeError, ValueError):
        return str(value)


def _pct(value, decimals=1):
    if value is None:
        return "—"
    try:
        return f"{float(value):.{decimals}f}%"
    except (TypeError, ValueError):
        return str(value)


def _set_cell_bg(cell, hex_color):
    """Aplica color de fondo a una celda (workaround python-docx)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_heading(doc, text, level=1, color="primary"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = COLORS[color]
    sizes = {1: 18, 2: 14, 3: 12}
    run.font.size = Pt(sizes.get(level, 11))
    run.font.name = "Arial"
    return p


def _add_paragraph(doc, text, bold=False, color=None, size=11, align=None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = COLORS[color]
    return p


def _prettify_header(header):
    """
    Convierte 'cpa_cop' → 'CPA COP', 'gasto_cop' → 'Gasto COP', etc.
    Capitaliza palabras simples ('nombre' → 'Nombre').
    Respeta headers que ya tienen mayúsculas/espacios.
    """
    if not isinstance(header, str):
        return str(header)
    # Reemplazos específicos comunes (snake_case y minúsculas)
    replacements = {
        "cpa_cop": "CPA COP",
        "gasto_cop": "Gasto COP",
        "presupuesto_cop": "Presupuesto COP",
        "roas": "ROAS",
        "cpa": "CPA",
        "ctr": "CTR",
        "cpc": "CPC",
        "cpm": "CPM",
    }
    h_lower = header.lower()
    if h_lower in replacements:
        return replacements[h_lower]
    # Si ya tiene mayúsculas/espacios (ya formateado), no tocar
    if header != h_lower:
        return header
    # snake_case → Title Case
    if "_" in header:
        return header.replace("_", " ").title()
    # Palabra simple en minúsculas: capitalizar
    return header.capitalize()


def _add_table(doc, headers, rows, row_colors=None):
    """
    Agrega una tabla con headers + rows.
    row_colors: lista paralela a rows con un string de SHADE por fila, o None.
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _set_cell_bg(cell, SHADE["primary"])
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(_prettify_header(h))
        run.bold = True
        run.font.color.rgb = COLORS["white"]
        run.font.size = Pt(10)
        run.font.name = "Arial"

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[1 + r_idx]
        color_key = (row_colors[r_idx] if row_colors and r_idx < len(row_colors) else None)
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Arial"
            if color_key and color_key in SHADE:
                _set_cell_bg(cell, SHADE[color_key])

    return table


def _semaforo_color(veredicto):
    """Mapea veredicto a clave SHADE."""
    v = str(veredicto).upper() if veredicto else ""
    if "🟢" in v or "ESCALAR" in v or "CONTINUAR" in v or "EXCELENTE" in v or "BUENO" in v:
        return "lightG"
    if "🟡" in v or "REVISAR" in v or "MARGINAL" in v or "ACEPTABLE" in v:
        return "lightY"
    if "🔴" in v or "PAUSAR" in v or "PIERDE" in v:
        return "lightR"
    if "⚫" in v or "DESCARTAR" in v:
        return "lightR"
    return None


# ============================================================
# SECCIONES DEL REPORTE
# ============================================================

def _section_portada(doc, cfg):
    _add_paragraph(doc, "ANÁLISIS DE CAMPAÑAS META ADS",
                   bold=True, color="primary", size=22, align="center")
    _add_paragraph(doc, cfg.get("producto", "—").upper(),
                   bold=True, color="accent", size=18, align="center")
    _add_paragraph(doc, "Diagnóstico ejecutable · Golden Group", size=11, align="center")

    doc.add_paragraph()

    # Datos clave
    periodo = cfg.get("periodo", "—")
    totales = cfg.get("totales", {})
    gasto = _cop(totales.get("gasto_cop"))
    compras = totales.get("compras", "—")
    cpa = _cop(totales.get("cpa_promedio"))

    _add_paragraph(doc, f"Periodo: {periodo}", bold=True, size=12, align="center")
    _add_paragraph(doc, f"Inversión: {gasto} COP · Compras: {compras} · CPA promedio: {cpa} COP",
                   size=11, align="center")

    # Aviso de conversión USD→COP si aplica
    if cfg.get("moneda_original") == "USD":
        trm = cfg.get("trm_aplicada")
        fuente = cfg.get("trm_fuente", "—")
        fuente_txt = {
            "web_search": "tasa real del día (consultada en internet)",
            "fallback_4000": "tasa de referencia conservadora $4,000 COP/USD (sin acceso a internet)",
            "usuario": "tasa indicada por el usuario",
        }.get(fuente, fuente)
        _add_paragraph(doc,
                       f"⚠ Reporte original en USD. Convertido a COP usando TRM ${trm:,.0f} COP/USD ({fuente_txt}). Todos los valores del documento están en COP.",
                       size=10, color="accent", align="center")

    doc.add_page_break()


def _section_unit_economics(doc, cfg):
    _add_heading(doc, "1. Unit Economics", level=1)

    ue = cfg.get("unit_economics") or {}
    if not ue:
        _add_paragraph(doc, "Sin datos de unit economics provistos.", color="accent")
        return

    flujo = ue.get("flujo_100_ordenes", {})
    pyl = ue.get("pyl", {})

    _add_heading(doc, "1.1 Flujo real sobre 100 órdenes Meta", level=2)
    _add_table(doc,
        headers=["Estado", "Cantidad", "% del total"],
        rows=[
            ["Canceladas",  flujo.get("canceladas", "—"), _pct(flujo.get("canceladas", 0))],
            ["Enviadas",    flujo.get("enviadas", "—"),   _pct(flujo.get("enviadas", 0))],
            ["Devueltas",   flujo.get("devueltas", "—"),  _pct((flujo.get("devueltas", 0) / 100) * 100)],
            ["Entregadas",  flujo.get("entregadas", "—"), _pct(flujo.get("tasa_entrega_real_pct", 0))],
        ])

    _add_heading(doc, "1.2 P&L sobre 100 órdenes", level=2)
    _add_table(doc,
        headers=["Concepto", "Valor COP"],
        rows=[
            ["Ingresos",                  _cop(pyl.get("ingresos"))],
            ["Costo de producto",         _cop(pyl.get("costo_producto_total"))],
            ["Fletes totales",            _cop(pyl.get("fletes_total"))],
            ["Margen pre-pauta",          _cop(pyl.get("margen_pre_pauta"))],
            ["Margen % sobre ingresos",   _pct(pyl.get("margen_pre_pauta_pct"))],
        ])

    # CPA breakeven destacado
    doc.add_paragraph()
    _add_paragraph(doc, f">>> CPA BREAKEVEN: {_cop(ue.get('cpa_breakeven'))} COP",
                   bold=True, color="green", size=18, align="center")
    _add_paragraph(doc, f">>> ROAS breakeven: {ue.get('roas_breakeven', '—')}x",
                   bold=True, color="green", size=14, align="center")
    doc.add_paragraph()

    # CPA objetivo
    _add_heading(doc, "1.3 CPA objetivo según margen neto deseado", level=2)
    obj = ue.get("cpa_por_margen_objetivo", {})
    if obj:
        rows = []
        for key, val in obj.items():
            pct_label = key.replace("margen_", "").replace("pct", "%")
            rows.append([
                f"Margen {pct_label}",
                _cop(val.get("cpa_max")),
                f"{val.get('roas_objetivo', '—')}x",
                _cop(val.get("ganancia_por_100_ordenes")),
            ])
        _add_table(doc,
            headers=["Margen deseado", "CPA máximo COP", "ROAS objetivo", "Ganancia / 100 órdenes"],
            rows=rows,
            row_colors=["lightG"] * len(rows))


def _section_pyl_historico(doc, cfg):
    pyl = cfg.get("pyl_historico") or {}
    if not pyl:
        return
    _add_heading(doc, "2. P&L histórico real de la cuenta", level=1)

    rentable = pyl.get("rentable")
    icon = "🟢" if rentable else "🔴"
    veredicto = "GANÓ PLATA" if rentable else "PERDIÓ PLATA"
    color = "green" if rentable else "accent"
    _add_paragraph(doc, f"{icon} La cuenta {veredicto} en el periodo analizado",
                   bold=True, color=color, size=14)

    _add_table(doc,
        headers=["Concepto", "Valor"],
        rows=[
            ["Gasto en pauta",          _cop(pyl.get("gasto_pauta"))],
            ["Compras Meta",            pyl.get("compras_meta", "—")],
            ["CPA promedio real",       _cop(pyl.get("cpa_promedio"))],
            ["Entregas reales",         pyl.get("entregadas_reales", "—")],
            ["Ingresos reales",         _cop(pyl.get("ingresos_reales"))],
            ["Margen pre-pauta",        _cop(pyl.get("margen_pre_pauta"))],
            ["MARGEN NETO FINAL",       _cop(pyl.get("margen_neto_final"))],
            ["Margen %",                _pct(pyl.get("margen_pct"))],
        ])


def _section_ranking(doc, cfg, key, titulo, num):
    items = cfg.get(key) or []
    if not items:
        return
    _add_heading(doc, f"{num}. {titulo}", level=1)
    # items: lista de dicts con nombre/gasto/compras/cpa/roas/veredicto
    headers = list(items[0].keys())
    rows = [[item.get(h, "—") for h in headers] for item in items]
    row_colors = [_semaforo_color(item.get("veredicto")) for item in items]
    _add_table(doc, headers=headers, rows=rows, row_colors=row_colors)


def _section_demografia(doc, cfg, num):
    demo = cfg.get("demografia_tabla") or []
    if not demo:
        return
    _add_heading(doc, f"{num}. Demografía (Edad × Sexo)", level=1)
    headers = list(demo[0].keys())
    rows = [[d.get(h, "—") for h in headers] for d in demo]
    row_colors = [_semaforo_color(d.get("veredicto")) for d in demo]
    _add_table(doc, headers=headers, rows=rows, row_colors=row_colors)


def _section_ubicaciones(doc, cfg, num):
    ubic = cfg.get("ubicaciones_tabla") or []
    if not ubic:
        return
    _add_heading(doc, f"{num}. Plataformas y ubicaciones", level=1)
    headers = list(ubic[0].keys())
    rows = [[u.get(h, "—") for h in headers] for u in ubic]
    row_colors = [_semaforo_color(u.get("veredicto")) for u in ubic]
    _add_table(doc, headers=headers, rows=rows, row_colors=row_colors)


def _section_landings(doc, cfg, num):
    ldg = cfg.get("landings_tabla") or []
    if not ldg:
        return
    _add_heading(doc, f"{num}. Landings / Páginas de destino", level=1)
    headers = list(ldg[0].keys())
    rows = [[l.get(h, "—") for h in headers] for l in ldg]
    _add_table(doc, headers=headers, rows=rows)


def _section_embudo(doc, cfg, num):
    em = cfg.get("embudo") or {}
    if not em:
        return
    _add_heading(doc, f"{num}. Embudo de conversión", level=1)
    rows = []
    for etapa, datos in em.items():
        if isinstance(datos, dict):
            rows.append([etapa, datos.get("tasa", "—"), datos.get("benchmark", "—"), datos.get("veredicto", "—")])
        else:
            rows.append([etapa, datos, "—", "—"])
    _add_table(doc,
        headers=["Etapa", "Tasa actual", "Benchmark", "Veredicto"],
        rows=rows,
        row_colors=[_semaforo_color(r[3]) for r in rows])


def _section_plan_accion(doc, cfg, num):
    plan = cfg.get("plan_accion") or {}
    if not plan:
        return
    _add_heading(doc, f"{num}. Plan de acción", level=1)

    if plan.get("pausar"):
        _add_heading(doc, "A. Qué pausar HOY", level=2, color="accent")
        for item in plan["pausar"]:
            if isinstance(item, dict):
                _add_paragraph(doc,
                    f"• {item.get('nombre', '—')} — {item.get('razon', '')} (gasto: {_cop(item.get('gasto_cop'))})")
            else:
                _add_paragraph(doc, f"• {item}")

    if plan.get("escalar"):
        _add_heading(doc, "B. Qué continuar y ESCALAR", level=2, color="green")
        for item in plan["escalar"]:
            if isinstance(item, dict):
                _add_paragraph(doc,
                    f"• {item.get('nombre', '—')} — CPA actual {_cop(item.get('cpa_cop'))}, presupuesto sugerido {_cop(item.get('presupuesto_cop'))}/día")
            else:
                _add_paragraph(doc, f"• {item}")

    if plan.get("campañas_nuevas"):
        _add_heading(doc, "C. Configuración de campañas NUEVAS", level=2, color="primary")
        for c in plan["campañas_nuevas"]:
            if isinstance(c, dict):
                _add_paragraph(doc, c.get("nombre", "—"), bold=True)
                for k, v in c.items():
                    if k != "nombre":
                        _add_paragraph(doc, f"   {k}: {v}", size=10)
            else:
                _add_paragraph(doc, str(c))

    if plan.get("replicar_mismo_bm"):
        _add_heading(doc, "D. Replicar en otra cuenta del mismo BM", level=2)
        for item in plan["replicar_mismo_bm"]:
            _add_paragraph(doc, f"• {item}")

    if plan.get("replicar_bm_nuevo"):
        _add_heading(doc, "E. Replicar en BM nuevo (cuenta desde cero)", level=2)
        for item in plan["replicar_bm_nuevo"]:
            _add_paragraph(doc, f"• {item}")


def _section_sintesis(doc, cfg, num):
    sintesis = cfg.get("sintesis_ejecutiva") or []
    if not sintesis:
        return
    _add_heading(doc, f"{num}. Síntesis ejecutiva", level=1)
    for i, punto in enumerate(sintesis, 1):
        _add_paragraph(doc, f"{i}. {punto}")


def _section_produccion_creativa(doc, cfg, num):
    pc = cfg.get("produccion_creativa") or {}
    if not pc:
        return
    _add_heading(doc, f"{num}. Producción creativa recomendada", level=1)
    for k, v in pc.items():
        _add_paragraph(doc, f"{k}: {v}")


# ============================================================
# API PRINCIPAL
# ============================================================

def generate_word_report(cfg, output_path):
    """
    Genera el .docx a partir del config_dict.
    Cada sección se omite automáticamente si su data falta.

    Returns: ruta absoluta al archivo generado.
    """
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Fuente base
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    _section_portada(doc, cfg)
    _section_unit_economics(doc, cfg)
    _section_pyl_historico(doc, cfg)

    num = 3
    if cfg.get("campañas_ranking"):
        _section_ranking(doc, cfg, "campañas_ranking", "Ranking de campañas", num); num += 1
    if cfg.get("creativos_ranking"):
        _section_ranking(doc, cfg, "creativos_ranking", "Ranking de creativos", num); num += 1
    if cfg.get("demografia_tabla"):
        _section_demografia(doc, cfg, num); num += 1
    if cfg.get("ubicaciones_tabla"):
        _section_ubicaciones(doc, cfg, num); num += 1
    if cfg.get("landings_tabla"):
        _section_landings(doc, cfg, num); num += 1
    if cfg.get("embudo"):
        _section_embudo(doc, cfg, num); num += 1
    if cfg.get("plan_accion"):
        _section_plan_accion(doc, cfg, num); num += 1
    if cfg.get("sintesis_ejecutiva"):
        _section_sintesis(doc, cfg, num); num += 1
    if cfg.get("produccion_creativa"):
        _section_produccion_creativa(doc, cfg, num); num += 1

    # Footer
    doc.add_paragraph()
    _add_paragraph(doc, f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')} · Golden Group · Análisis estándar Meta Ads",
                   size=9, color="gray", align="center")

    output_path = str(Path(output_path).resolve())
    doc.save(output_path)
    return output_path


def build_output_filename(producto, fecha_analisis=None):
    """
    Construye el nombre estándar: 'Analisis de <Producto> <YYYY-MM-DD>.docx'

    - No usa 'versión corregida', 'v2', 'final' nunca.
    - Fecha = día en que se corre el análisis (no la del informe).
    """
    if fecha_analisis is None:
        fecha_analisis = datetime.now().strftime("%Y-%m-%d")
    # Sanitizar el nombre del producto para filesystem
    producto_clean = "".join(c if c.isalnum() or c in " -_." else "_" for c in producto).strip()
    return f"Analisis de {producto_clean} {fecha_analisis}.docx"


# ============================================================
# CLI
# ============================================================

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Uso: python generate_report.py <config.json> <output.docx>")
        sys.exit(1)

    config_path = sys.argv[1]
    output_path = sys.argv[2]

    with open(config_path, "r", encoding="utf-8") as f:
        cfg = json.load(f)

    result = generate_word_report(cfg, output_path)
    print(f"✅ Word generado: {result}")
